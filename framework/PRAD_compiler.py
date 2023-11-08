import docx
import pandas as pd

class dataCompiler:
    def __init__(self, path, arquivo):
        self.path = path
        self.arquivo = arquivo
        self.colunas = ['Arquivo']
        self.tecnicas = ['Plantio', 'Adensamento', 'Jardim', 'Regeneração', 'Enriquecimento', 'Remoção', 'Adicionalidade']
        self.raw_data = list()
        self.filt_data = list()

    def iter_unique_cells(self, row):
        prior_tc = None
        for cell in row.cells:
            this_tc = cell._tc
            if this_tc is prior_tc:
                continue
            prior_tc = this_tc
            yield cell

    def readPRAD(self):
        curr_doc = docx.Document(self.path + self.arquivo) #lê o aqruivo docx
        raw_entries = list() #lista que registra todos os dados da tabela 5 (técnicas)
        for tabela in curr_doc.tables: #itera entre as tabelas existentes no doc
            dadosProc = False #flag que registra quando inicia a tabela 5 (e quando termina)
            for linha in tabela.rows: #itera entre as linhas da tabela da iteração atual
                for cel in self.iter_unique_cells(linha): #identifica entradas ñ repetidas
                    for paragrafo in cel.paragraphs: #para cada "nova palavra"
                        if paragrafo.text == '5. PROCEDIMENTOS DA RESTAURAÇÃO ':
                            dadosProc = True
                        elif paragrafo.text == '6. LISTA DE ESPÉCIES, QUANTIDADE E DISTRIBUIÇÃO POR POLÍGONO':
                            dadosProc = False

                        if dadosProc: #se a leitura é referente a tabela 5
                            if len(self.colunas) < 7: #monta a lista de colunas do dataframe
                                if paragrafo.text == '5. PROCEDIMENTOS DA RESTAURAÇÃO ': #nome da tab 5 ñ entra
                                    continue
                                elif len(self.colunas) > 0: #se a tabela ñ estiver vazia
                                    if self.colunas[-1] == paragrafo.text: #se repetir colunas devido ao merge do word
                                        continue
                                    else: #tudo ok, preenche a lista com a frase
                                        self.colunas.append(paragrafo.text)
                                else: #tudo ok, preenche a lista com a frase
                                    self.colunas.append(paragrafo.text)
                            else:
                                if paragrafo.text not in self.colunas: #checa se frase é uma entrada da tabela
                                    raw_entries.append(paragrafo.text)

        return raw_entries

    def filtDup(self, raw_entries):
        for data in range(len(raw_entries)): #itera no numero de entradas registradas originalmente
            if data%(len(self.colunas)-1) == 0: #tabela possui 6 colunas (a cada 6 valores é uma linha nova)
                if data < len(raw_entries) and raw_entries[data] == raw_entries[data-1]: #checa se está duplicada a técnica apenas
                    raw_entries.pop(data) #retira se verdade

        idx_tecs = list() #registra ocorrência dos valores de técnica
        for idx, frase in enumerate(raw_entries): #itera entre os valores originais das entradas
            for tec in self.tecnicas: #para as técnicas que existem
                if tec in frase: #achou a tecnica, registra em qual índice
                    idx_tecs.append(idx)
                    break

        sep_entries = list() #separa as entradas originais dados os índices anteriores
        for new_line in range(len(idx_tecs)): #itera entre os índices de cada nova linha
            if new_line < len(idx_tecs) - 1: #até o penult. índice tem que verificar o seguinte
                sep_entries.append(raw_entries[int(new_line*(len(self.colunas)-1)):int((new_line+1)*(len(self.colunas)-1))])
            else: #não existe próximo no último, vai até o ultimo elemento então
                sep_entries.append(raw_entries[idx_tecs[-1]:])

        filt_entries = list() #registra a entradas agora filtradas para duplicadas em todos elementos
        for linhas in range(len(sep_entries)):
            idx_drop = list() #índice de quais repetidas devem ser retiradas
            palavra_atual = '' #registra qual está checando para duplicadas
            for cols in range(len(sep_entries[linhas])): #para cada linha verifica elemento por elemento
                if sep_entries[linhas][cols] != '--': #'--' pode ser repetido com frequência
                    if palavra_atual == sep_entries[linhas][cols]: #registrou uma repetição
                        idx_drop.append(cols)
                    else:
                        palavra_atual = sep_entries[linhas][cols] #comuta a palavra pq nova

            sep_entries[linhas] = sep_entries[linhas][:] #ñ lembro pq coloquei isso, verificar dps
            filt_line = [sep_entries[linhas][cols] for cols in range(len(sep_entries[linhas]))
                         if cols not in idx_drop] #dados os índices das repetições, cria uma lista filtrada
            filt_entries.append(filt_line)

        return filt_entries

    def filtForm(self, filt_entries):
        for correc_i in range(len(filt_entries)): #itera entre as entrada filtradas
            count = 0 #registra o número de operações de append para cada linha filtrada
            if len(filt_entries[correc_i]) < (len(self.colunas)-1): #se a linha não estiver preenchida com as 6 colunas
                while(len(filt_entries[correc_i]) < (len(self.colunas)-1)): #preenche até as 6 colunas completas com dados
                    if filt_entries[correc_i][-1] == '--': #caso especial de '--' no final
                        if filt_entries[correc_i][0].isnumeric():
                            filt_entries[correc_i].pop(0) #primeira coluna nunca é número
                        else:
                            if correc_i < len(filt_entries) - 1: #operações que necessitam checar um indice seg.
                                if filt_entries[correc_i + 1][0].isnumeric(): #checa se entrada seguinte é numeral
                                    if count == 0: #caso especial sem manipulações de pop() e append() correndo
                                        filt_entries[correc_i].append(filt_entries[correc_i + 1][0])
                                    else: #caso especial de manipulações de pop() e append() além do índice [0]
                                        filt_entries[correc_i].append(filt_entries[correc_i + 1][count])
                                        count += 1 #aumenta contador
                                else:
                                    filt_entries[correc_i].append('--') #caso de '--' em sequência
                    else:
                        if not filt_entries[correc_i][0].isnumeric(): #ñ tem 6 entradas ainda mas a primeira é numeral
                            if correc_i < len(filt_entries) - 1: #operações que necessitam checar um indice seg.
                                filt_entries[correc_i].append(filt_entries[correc_i + 1][count])
                                #add valor da linha seguinte relativo ao índice count
                                count += 1 #aumenta contador
                        else:
                            filt_entries[correc_i].pop(0) #se for último índice e o primeiro é numeral, retira
            else: #6 entradas preenchidas na linha
                if filt_entries[correc_i][0].isnumeric(): #se estiver completo mas [0] ñ é tecnica
                    filt_entries[correc_i].pop(0) #retira o índice
                    #seguir: mesmo processo para o caso em que está com menos de 6 entradas
                    while(len(filt_entries[correc_i]) < (len(self.colunas)-1)):
                        if filt_entries[correc_i][-1] == '--':
                            if filt_entries[correc_i][0].isnumeric():
                                filt_entries[correc_i].pop(0)
                            else:
                                if correc_i < len(filt_entries) - 1:
                                    if filt_entries[correc_i + 1][0].isnumeric():
                                        filt_entries[correc_i].append(filt_entries[correc_i + 1][0])
                                    else:
                                        filt_entries[correc_i].append('--')
                        else:
                            if not filt_entries[correc_i][0].isnumeric():
                                if correc_i < len(filt_entries) - 1:
                                    filt_entries[correc_i].append(filt_entries[correc_i + 1][0])
                elif filt_entries[correc_i][1].isnumeric():
                    #caso especial do merge jogar dois valores para início da linha seguinte
                    filt_entries[correc_i] = filt_entries[correc_i][2:] #desconsidera esses dois
                    if correc_i < len(filt_entries) - 1: #dois valoes no fim da linha anterior
                        filt_entries[correc_i].append(filt_entries[correc_i + 1][0])
                        filt_entries[correc_i].append(filt_entries[correc_i + 1][1])
                else: #nenhum dos cenários, apenas segue para próx. iteração
                    continue

        return filt_entries

    def runFilters(self, raw_entries):
        f_dup = self.filtDup(raw_entries)
        f_form = self.filtForm(f_dup)
        f_final = list() #última verificação de forma qto ao formato
        for dup in range(len(f_form)):
            m_form = f_form[dup][0] #primeira coluna tem que ser técnica
            m_form = m_form.split() #divide em palavras
            popFlag = False #flag para registrar algum possível erro de formato
            if len(m_form) > 0:
                if m_form[0] in self.tecnicas: #se a primeira palavra de técnica é realmente técnica
                    idx = 1
                    for weird_occ in f_form[dup][1:]: #casoespecial meio da linha existe uma entrada de técnica
                        check = weird_occ.split() #divide em palavras
                        if len(check) > 0:
                            if check[0] in self.tecnicas:
                                if idx == len(f_form[dup]) - 1:
                                    popFlag = False
                                    f_form[dup][idx] = '--'
                                else:
                                    popFlag = True #achou erro!
                        idx += 1
                else: #primeiro item não é técnica
                    if dup > 0 and dup < (len(f_form) - 1):
                        if f_form[dup][-1].isnumeric():
                            f_form[dup] = f_form[dup][1:]
                            if dup != (len(f_form) - 2):
                                f_form[dup].append(f_form[dup + 1][0])
                            else:
                                f_form[dup].append('--')
                            popFlag = False
                        else:
                            splice = f_form[dup][-1].split()
                            if splice[0] in self.tecnicas:
                                f_form[dup] = f_form[dup][1:-1]
                                f_form[dup].append('--')
                                f_form[dup].append('--')
                                popFlag = False
                            else:
                                if splice[0] == '--':
                                    f_form[dup] = f_form[dup][1:-1]
                                    f_form[dup].append('--')
                                    f_form[dup].append('--')
                                    popFlag = False
                                else:
                                    popFlag = True
                    else:
                        popFlag = True

            if not popFlag: #caso flag seja False (sem erros)
                f_final.append(f_form[dup])
        return f_final

    def convDec(self, filt_data):
        for linha in filt_data: # para todas as linhas com entradas
            for i in range(len(linha)): #iteração elemento por elemento
                linha[i] = linha[i].replace(',', '.') #decimal separado por . para converter para int ou float
                if linha[i] == '--' or linha[i] == '-' or linha[i] == '---' or linha[i] == '----' or linha[i] == '':
                    linha[i] = 0

    def runCompile(self):
        self.raw_data = self.readPRAD() #lê .docx
        self.filt_data = self.runFilters(self.raw_data) #roda os filtros
        self.convDec(self.filt_data)
        PRAD_df = pd.DataFrame(columns=self.colunas) #cria o dataframe do PRAD atual
        for ents in range(len(self.filt_data)): #itera para cada linha dos dados filtrados
            entry = {self.colunas[0]: self.arquivo, #nome do arquivo
                     self.colunas[1]: self.filt_data[ents][0], #tecnica de restauração
                     self.colunas[2]: float(self.filt_data[ents][1]), #área (ha)
                     self.colunas[3]: self.filt_data[ents][2], #n polígono
                     self.colunas[4]: int(self.filt_data[ents][3]), #n mudas nativas
                     self.colunas[5]: int(self.filt_data[ents][4]), #n ervas
                     self.colunas[6]: float(self.filt_data[ents][5])} #metragem cercas
            PRAD_df.loc[len(PRAD_df)] = entry

        return PRAD_df

